# =============================================================================
# RetrievalLab — corpus/loaders/text_loader.py
# =============================================================================
# PURPOSE : Concrete loaders for plain text, Markdown, DOCX, and HTML files.
#           All implement BaseLoader; all return ParsedDocument objects.
#
# LOADERS DEFINED HERE:
#   TextLoader     — .txt, .text files (raw UTF-8)
#   MarkdownLoader — .md, .markdown files (strips frontmatter + syntax)
#   DocxLoader     — .docx files via python-docx (paragraphs + tables)
#   HTMLLoader     — .html, .htm files via BeautifulSoup4 (main content only)
#
# DESIGN NOTE:
#   These loaders are intentionally simple — they handle 80% of real-world
#   files with minimal dependencies. Exotic formats (EPUB, RTF, PPTX) can
#   be added as separate loaders without touching this file.
#
# INPUT  : File path for each respective format
# OUTPUT : ParsedDocument (same contract as PDFLoader)
#
# REMAINING AFTER THIS FILE:
#   Chunks are not created here — that's ChunkEngine's job.
#   These loaders only extract text, metadata, and tables.
# =============================================================================

from __future__ import annotations

import re
from pathlib import Path

import structlog

from corpus.loaders.base_loader import BaseLoader, ParsedDocument

logger = structlog.get_logger(__name__)


# =============================================================================
# TextLoader
# =============================================================================


class TextLoader(BaseLoader):
    """
    Loader for plain text files (.txt, .text, .log, .csv as text).

    No structural parsing — treats entire file content as a single document.
    Handles encoding detection with UTF-8 fallback to latin-1.

    Args:
        encoding: Force a specific encoding. Default: auto-detect (UTF-8 → latin-1).

    Input : .txt / .text file (any size up to MAX_FILE_SIZE_BYTES)
    Output: ParsedDocument with full file text, no tables, minimal metadata
    """

    SUPPORTED_EXTENSIONS = {".txt", ".text", ".log"}

    def __init__(self, encoding: str | None = None, **kwargs) -> None:
        super().__init__(**kwargs)
        self.encoding = encoding

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def _parse(self, path: Path) -> ParsedDocument:
        # Try specified encoding, then UTF-8, then latin-1 as last resort
        encodings_to_try = [self.encoding] if self.encoding else ["utf-8", "latin-1"]
        text = ""
        used_encoding = "utf-8"

        for enc in encodings_to_try:
            try:
                text = path.read_text(encoding=enc)
                used_encoding = enc
                break
            except UnicodeDecodeError:
                continue

        return ParsedDocument(
            text=text,
            source=str(path),
            metadata={
                "source_format": "text",
                "encoding": used_encoding,
                "file_name": path.name,
                "file_size_bytes": path.stat().st_size,
            },
        )


# =============================================================================
# MarkdownLoader
# =============================================================================


class MarkdownLoader(BaseLoader):
    """
    Loader for Markdown files (.md, .markdown).

    Strips:
    - YAML/TOML frontmatter (---...--- or +++...+++)
    - Code fence blocks (```...```) — keeps content, removes fences
    - Markdown syntax characters (#, *, _, ~~, >, |)

    Preserves:
    - Heading text (useful for section-aware chunking)
    - Paragraph structure (blank lines between sections)

    Input : .md / .markdown file
    Output: ParsedDocument with cleaned plain text, section headers in metadata
    """

    SUPPORTED_EXTENSIONS = {".md", ".markdown", ".mdx"}

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def _parse(self, path: Path) -> ParsedDocument:
        try:
            raw = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raw = path.read_text(encoding="latin-1")

        # Strip YAML frontmatter (--- ... ---)
        raw = re.sub(r"^---\s*\n.*?\n---\s*\n", "", raw, flags=re.DOTALL)
        # Strip TOML frontmatter (+++ ... +++)
        raw = re.sub(r"^\+\+\+\s*\n.*?\n\+\+\+\s*\n", "", raw, flags=re.DOTALL)

        # Extract section headers for metadata
        headers = re.findall(r"^#{1,6}\s+(.+)$", raw, flags=re.MULTILINE)

        # Remove code fences but keep code content (useful for tech corpora)
        raw = re.sub(r"```\w*\n?", "", raw)

        # Strip markdown syntax
        raw = re.sub(r"!\[.*?\]\(.*?\)", "", raw)  # images
        raw = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", raw)  # links → link text
        raw = re.sub(r"[*_]{1,2}([^*_]+)[*_]{1,2}", r"\1", raw)  # bold/italic
        raw = re.sub(r"~~([^~]+)~~", r"\1", raw)  # strikethrough
        raw = re.sub(r"^>\s+", "", raw, flags=re.MULTILINE)  # blockquotes
        raw = re.sub(r"^[-*+]\s+", "", raw, flags=re.MULTILINE)  # list bullets
        raw = re.sub(r"^\d+\.\s+", "", raw, flags=re.MULTILINE)  # numbered lists
        raw = re.sub(r"^\|.*\|$", "", raw, flags=re.MULTILINE)  # table rows
        raw = re.sub(r"^#{1,6}\s+", "", raw, flags=re.MULTILINE)  # headers

        return ParsedDocument(
            text=raw,
            source=str(path),
            metadata={
                "source_format": "markdown",
                "file_name": path.name,
                "section_headers": headers[:20],  # first 20 headers for preview
                "header_count": len(headers),
            },
        )


# =============================================================================
# DocxLoader
# =============================================================================


class DocxLoader(BaseLoader):
    """
    Loader for Microsoft Word documents (.docx).

    Uses python-docx to extract:
    - Paragraph text in document reading order
    - Table cell contents (linearized row by row)
    - Core document metadata (title, author, created date)

    Skips: embedded images, SmartArt, footnotes (Day 2+ enhancement).

    Args:
        include_tables: Whether to include table text in the main text body.
                        Default True. Tables are always stored in doc.tables too.

    Input : .docx file
    Output: ParsedDocument with full paragraph + table text and metadata

    NOTE: .doc (legacy binary format) is NOT supported. Convert to .docx first
          using LibreOffice: `soffice --headless --convert-to docx file.doc`
    """

    SUPPORTED_EXTENSIONS = {".docx"}

    def __init__(self, include_tables: bool = True, **kwargs) -> None:
        super().__init__(**kwargs)
        self.include_tables = include_tables

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def _parse(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        doc = DocxDocument(str(path))

        # ── Extract paragraph text ─────────────────────────────────────────
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # ── Extract tables ─────────────────────────────────────────────────
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            table_data: list[list[str]] = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                # Deduplicate merged cells (python-docx repeats merged cells)
                deduplicated = []
                for i, cell_text in enumerate(row_data):
                    if i == 0 or cell_text != row_data[i - 1]:
                        deduplicated.append(cell_text)
                table_data.append(deduplicated)
            if table_data:
                tables.append(table_data)

                # Also include table as text if requested
                if self.include_tables:
                    table_text = "\n".join(" | ".join(row) for row in table_data)
                    paragraphs.append(f"\n[TABLE]\n{table_text}\n[/TABLE]")

        # ── Extract core metadata ──────────────────────────────────────────
        core = doc.core_properties
        metadata = {
            "source_format": "docx",
            "file_name": path.name,
            "title": core.title or "",
            "author": core.author or "",
            "created": str(core.created) if core.created else "",
            "modified": str(core.modified) if core.modified else "",
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        }

        full_text = "\n\n".join(paragraphs)

        return ParsedDocument(
            text=full_text,
            source=str(path),
            metadata=metadata,
            tables=tables,
        )


# =============================================================================
# HTMLLoader
# =============================================================================


class HTMLLoader(BaseLoader):
    """
    Loader for HTML files and web pages.

    Uses BeautifulSoup4 to:
    - Extract main content text (removes nav, header, footer, scripts, ads)
    - Parse page title from <title> tag
    - Extract meta description
    - Follow main content heuristics (<main>, <article>, <div id="content">)

    Intended for: web-scraped corpora, documentation sites, Wikipedia dumps.

    Input : .html / .htm file
    Output: ParsedDocument with cleaned body text and page metadata
    """

    SUPPORTED_EXTENSIONS = {".html", ".htm", ".xhtml"}

    # Tags to strip entirely (content not useful for retrieval)
    REMOVE_TAGS = {
        "script",
        "style",
        "nav",
        "footer",
        "header",
        "aside",
        "advertisement",
        "noscript",
        "iframe",
    }

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def _parse(self, path: Path) -> ParsedDocument:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("BeautifulSoup4 required: pip install beautifulsoup4 lxml")

        try:
            raw_html = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raw_html = path.read_text(encoding="latin-1")

        soup = BeautifulSoup(raw_html, "lxml")

        # ── Remove noise tags ──────────────────────────────────────────────
        for tag in soup.find_all(list(self.REMOVE_TAGS)):
            tag.decompose()

        # ── Extract metadata ───────────────────────────────────────────────
        title = soup.title.string.strip() if soup.title else ""
        meta_desc = ""
        if soup.find("meta", attrs={"name": "description"}):
            meta_desc = soup.find("meta", attrs={"name": "description"}).get("content", "")

        # ── Find main content container ────────────────────────────────────
        # Priority: <main> → <article> → <div id="content"> → <body>
        content_container = (
            soup.find("main")
            or soup.find("article")
            or soup.find("div", id=re.compile(r"content|main|body", re.I))
            or soup.find("body")
            or soup
        )

        # ── Extract text ───────────────────────────────────────────────────
        text = content_container.get_text(separator="\n", strip=True)

        # Collapse repeated newlines
        text = re.sub(r"\n{3,}", "\n\n", text)

        return ParsedDocument(
            text=text,
            source=str(path),
            metadata={
                "source_format": "html",
                "file_name": path.name,
                "title": title,
                "meta_description": meta_desc,
            },
        )
